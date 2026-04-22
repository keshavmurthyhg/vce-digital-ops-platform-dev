import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
import tempfile
import zipfile
import os
from datetime import datetime
from docx import Document


SECTION_HEADERS = [
    "New Parts",
    "Revised Parts",
    "New Documents",
    "Revised Documents",
    "Non Production / Obsolete / RG4 Parts",
    "Structure Changes",
    "Line Number Changes",
    "Associated sBOMs",
    "Serial No. / Effectivity"
]


def normalize(val):
    if pd.isna(val):
        return ""
    return str(val).strip()


# =========================
# LOAD DATA
# =========================
def compare_excels(file1, file2):
    df1 = pd.read_excel(file1, header=None)
    df2 = pd.read_excel(file2, header=None)

    df1 = df1.reset_index(drop=True)
    df2 = df2.reset_index(drop=True)

    return df1, df2


# =========================
# SECTION DETECTION
# =========================
def extract_sections(df):
    sections = {}
    current_section = None

    for idx, row in df.iterrows():
        first_val = normalize(row.iloc[0])

        if first_val in SECTION_HEADERS:
            current_section = first_val
            sections[current_section] = []

        elif current_section:
            sections[current_section].append(idx)

    return sections


# =========================
# FINAL DIFF LOGIC
# =========================
def section_diff_logic(df1, df2):
    sections1 = extract_sections(df1)
    sections2 = extract_sections(df2)

    diff_mask1 = pd.DataFrame("", index=df1.index, columns=df1.columns)
    diff_mask2 = pd.DataFrame("", index=df2.index, columns=df2.columns)

    summary = {}
    total_diff = 0
    removed_rows_data = []

    for section in SECTION_HEADERS:
        rows1 = sections1.get(section, [])
        rows2 = sections2.get(section, [])

        map1 = {}
        map2 = {}

        for i in rows1:
            val = normalize(df1.iloc[i, 0])
            if val and val != "Number" and val not in map1:
                map1[val] = i

        for i in rows2:
            val = normalize(df2.iloc[i, 0])
            if val and val != "Number" and val not in map2:
                map2[val] = i

        added = set(map2.keys()) - set(map1.keys())
        removed = set(map1.keys()) - set(map2.keys())
        common = set(map1.keys()) & set(map2.keys())

        modified = 0

        # 🟢 Added (file2)
        for key in added:
            idx = map2[key]
            diff_mask2.iloc[idx, :] = "added"

        # 🔴 Removed (file1)
        for key in removed:
            idx = map1[key]
            diff_mask1.iloc[idx, :] = "removed"

            removed_rows_data.append({
                "section": section,
                "number": key
            })

        # 🟡 Modified (column-level)
        for key in common:
            i1 = map1[key]
            i2 = map2[key]

            row1 = df1.iloc[i1].fillna("")
            row2 = df2.iloc[i2].fillna("")

            if not row1.equals(row2):
                for col in df2.columns:
                    if normalize(row1[col]) != normalize(row2[col]):
                        diff_mask2.loc[i2, col] = "modified"
                modified += 1

        summary[section] = {
            "added": len(added),
            "removed": len(removed),
            "modified": modified
        }

        total_diff += len(added) + len(removed) + modified

    return diff_mask1, diff_mask2, summary, total_diff, removed_rows_data


# =========================
# STYLE
# =========================
def style_dataframe(df, diff_mask):
    def highlight(row):
        styles = []
        for col in df.columns:
            val = diff_mask.loc[row.name, col]

            if val == "added":
                styles.append("background-color: #90EE90")  # green
            elif val == "modified":
                styles.append("background-color: #FFD700")  # yellow
            elif val == "removed":
                styles.append("background-color: #FF7F7F")  # red
            else:
                styles.append("")
        return styles

    return df.style.apply(highlight, axis=1)


# =========================
# EXCEL OUTPUT
# =========================
def generate_output(file1, file2):
    wb1 = load_workbook(file1)
    wb2 = load_workbook(file2)

    ws1 = wb1.active
    ws2 = wb2.active

    fill_added = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
    fill_modified = PatternFill(start_color="FFD700", end_color="FFD700", fill_type="solid")
    fill_removed = PatternFill(start_color="FF7F7F", end_color="FF7F7F", fill_type="solid")

    df1, df2 = compare_excels(file1, file2)
    diff_mask1, diff_mask2, summary, total_diff, removed_rows = section_diff_logic(df1, df2)

    # File2 (added + modified)
    for r in range(len(df2)):
        for c in range(len(df2.columns)):
            val = diff_mask2.iloc[r, c]
            if val == "added":
                ws2.cell(r + 1, c + 1).fill = fill_added
            elif val == "modified":
                ws2.cell(r + 1, c + 1).fill = fill_modified

    # File1 (removed)
    for r in range(len(df1)):
        for c in range(len(df1.columns)):
            if diff_mask1.iloc[r, c] == "removed":
                ws1.cell(r + 1, c + 1).fill = fill_removed

    base1 = os.path.splitext(file1.name)[0]
    base2 = os.path.splitext(file2.name)[0]

    name1 = f"{base1}_Removed.xlsx"
    name2 = f"{base2}_Highlighted.xlsx"

    temp_dir = tempfile.mkdtemp()

    path1 = os.path.join(temp_dir, name1)
    path2 = os.path.join(temp_dir, name2)

    wb1.save(path1)
    wb2.save(path2)

    # Word Report
    doc = Document()
    doc.add_heading("Excel Comparison Summary", 0)

    for section, data in summary.items():
        if any(data.values()):
            doc.add_paragraph(
                f"{section} → Added: {data['added']} | Removed: {data['removed']} | Modified: {data['modified']}"
            )

    doc.add_heading("Removed Items", level=1)
    for item in removed_rows:
        doc.add_paragraph(f"{item['section']} → {item['number']}")

    word_path = os.path.join(temp_dir, "Comparison_Report.docx")
    doc.save(word_path)

    # ZIP
    date_str = datetime.now().strftime("%d%b%Y")
    zip_name = f"Excel-Compare_{date_str}.zip"
    zip_path = os.path.join(temp_dir, zip_name)

    with zipfile.ZipFile(zip_path, "w") as z:
        z.write(path1, name1)
        z.write(path2, name2)
        z.write(word_path, "Comparison_Report.docx")

    return zip_path, zip_name
