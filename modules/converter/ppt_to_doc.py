import os
import tempfile
import subprocess

from docx import Document
from docx.shared import Inches

from modules.converter.ppt_metadata import extract_slide1_metadata


def render_ppt_to_images(ppt_path):
    """
    Convert entire PPT slides into PNG images.
    Preserves:
    - arrows
    - lines
    - text boxes
    - screenshots
    - annotations
    - shapes
    """

    temp_dir = tempfile.mkdtemp()

    subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "png",
            "--outdir",
            temp_dir,
            ppt_path
        ],
        check=True
    )

    images = []

    for file in sorted(os.listdir(temp_dir)):
        if file.endswith(".png"):
            images.append(
                os.path.join(temp_dir, file)
            )

    if not images:
        raise Exception(
            "No slide images generated from PPT"
        )

    return images


def add_header_table(doc, metadata):
    doc.add_heading(
        "INCIDENT REPORT",
        level=0
    )

    table = doc.add_table(
        rows=2,
        cols=2
    )
    table.style = "Table Grid"

    table.cell(0, 0).text = "Incident"
    table.cell(0, 1).text = metadata.get(
        "incident", ""
    )

    table.cell(1, 0).text = "Created Date"
    table.cell(1, 1).text = metadata.get(
        "created_date", ""
    )


def ppt_to_word(ppt_path, output_docx):
    metadata = extract_slide1_metadata(
        ppt_path
    )

    slide_images = render_ppt_to_images(
        ppt_path
    )

    doc = Document()

    add_header_table(
        doc,
        metadata
    )

    # skip slide1 because metadata already used
    for img in slide_images[1:]:
        doc.add_page_break()
        doc.add_picture(
            img,
            width=Inches(6.8)
        )

    doc.save(output_docx)

    return output_docx
