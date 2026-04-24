from docx import Document
from io import BytesIO
import os

from modules.converter.ppt_to_doc import ppt_to_word
from modules.report.doc_generator import generate_word_doc


def copy_content(source, target):
    """
    Safely copy paragraphs + tables + images
    """
    for block in source.element.body:
        target.element.body.append(block)


def generate_combined_report(ppt_path, snow_data, root, l2, res, images):

    # 🔹 Step 1: Generate SNOW doc
    snow_bytes = generate_word_doc(snow_data, root, l2, res, images)
    snow_doc = Document(BytesIO(snow_bytes))

    # 🔹 Step 2: Generate PPT doc
    ppt_temp = "temp_ppt.docx"
    ppt_to_word(ppt_path, ppt_temp)
    ppt_doc = Document(ppt_temp)

    # 🔹 Step 3: Create clean doc (IMPORTANT)
    final_doc = Document()

    # ❌ REMOVE default empty paragraph (fix blank page)
    final_doc._element.body.clear()

    # 🔹 Add SNOW content
    copy_content(snow_doc, final_doc)

    # 🔹 Add page break manually
    final_doc.add_page_break()

    # 🔹 Add PPT content (skip first heading to avoid duplicate title)
    skip_first = True

    for element in ppt_doc.element.body:
        # Skip duplicate "INCIDENT REPORT"
        if skip_first:
            skip_first = False
            continue

        final_doc.element.body.append(element)

    os.remove(ppt_temp)

    # 🔹 Save
    buffer = BytesIO()
    final_doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()
