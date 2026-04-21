from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import re
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from reportlab.platypus import KeepTogether

def clean_text(text):
    if not text:
        return ""
    return re.sub(r"How does the user want.*?\d+", "", str(text), flags=re.I).strip()

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def generate_word_doc(data, root, l2, res, images=None):
    doc = Document()
    
    title = doc.add_heading("INCIDENT REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.color.rgb = RGBColor(31, 78, 121)

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    def fill(r, c, key, val):
        h = table.rows[r].cells[c]
        v = table.rows[r].cells[c+1]
        h.text = key.upper()
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "D9D9D9")
        h._element.get_or_add_tcPr().append(shade)
        for run in h.paragraphs[0].runs:
            run.bold = True

        p = v.paragraphs[0]
        if key == "Incident":
            add_hyperlink(p, f"https://volvoitsm.service-now.com/nav_to.do?uri=incident.do?sysparm_query=number={val}", str(val))
        elif key == "Azure Bug":
            add_hyperlink(p, f"https://dev.azure.com/VolvoGroup-DVP/VCEWindchillPLM/_workitems/edit/{val}", str(val))
        elif key == "PTC Case":
            add_hyperlink(p, f"https://support.ptc.com/app/caseviewer/?case={val}", str(val))
        else:
            p.text = str(val)

    fill(0,0,"Incident",data.get("number"))
    fill(0,2,"Created By",data.get("created_by"))
    fill(1,0,"Azure Bug",data.get("azure_bug"))
    fill(1,2,"Created Date",data.get("created_date"))
    fill(2,0,"PTC Case",data.get("ptc_case"))
    fill(2,2,"Assigned To",data.get("assigned_to"))
    fill(3,0,"Priority",data.get("priority"))
    fill(3,2,"Resolved Date",data.get("resolved_date"))

    doc.add_paragraph("")
    t2 = doc.add_table(rows=2, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "SHORT DESCRIPTION"
    t2.rows[0].cells[1].text = "DESCRIPTION"
    t2.rows[1].cells[0].text = clean_text(data.get("short_description"))
    t2.rows[1].cells[1].text = clean_text(data.get("description"))

    for section in [("ROOT CAUSE", root, "root"), ("L2 ANALYSIS", l2, "l2"), ("RESOLUTION", res, "res")]:
        doc.add_heading(section[0], 1)
        doc.add_paragraph(section[1])
        if images and images.get(section[2]):
            doc.add_picture(images[section[2]], width=Inches(5))

    # FOOTER: 3-column layout using Tab Stops
    from docx.enum.text import WD_TAB_ALIGNMENT
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.clear()
    
    tab_stops = footer.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    
    run = footer.add_run(f"{data.get('number')}\tPage ")
    
    # Add dynamic page number
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    run2 = footer.add_run(f"\t{data.get('priority')}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf(data, root, l2, res, images=None):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("<b>INCIDENT REPORT</b>", styles["Title"]))
    elements.append(Spacer(1,10))

    def link(url, text):
        return Paragraph(f'<link href="{url}" color="blue">{text}</link>', styles["Normal"])

    def wrap(x):
        return Paragraph(str(x), styles["Normal"])

    # Table 1: Header Info
    table = Table([
        ["INCIDENT", link(f"https://volvoitsm.service-now.com/...", data.get('number')), "CREATED BY", wrap(data.get('created_by'))],
        ["AZURE BUG", wrap(data.get('azure_bug')), "CREATED DATE", wrap(data.get('created_date'))],
        ["PTC CASE", wrap(data.get('ptc_case')), "ASSIGNED TO", wrap(data.get('assigned_to'))],
        ["PRIORITY", wrap(data.get('priority')), "RESOLVED DATE", wrap(data.get('resolved_date'))]
    ], colWidths=[80, 185, 80, 185])
    
        table.setStyle(TableStyle([
        ('GRID',(0,0),(-1,-1),1,colors.black),
        ('BACKGROUND',(0,0),(0,-1),colors.lightgrey),
        ('BACKGROUND',(2,0),(2,-1),colors.lightgrey),
        ('FONTSIZE', (0,0), (-1,-1), 9),
    ]))
    elements.append(KeepTogether(table))
    elements.append(Spacer(1, 15))

    # Table 2: Descriptions with Wrapping
    desc_table = Table([
        [Paragraph("<b>SHORT DESCRIPTION</b>", styles["Normal"]), Paragraph("<b>DESCRIPTION</b>", styles["Normal"])],
        [wrap(clean_text(data.get("short_description"))), wrap(clean_text(data.get("description")))]
    ], colWidths=[180, 350])
                         
desc_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 1, colors.black),
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('VALIGN', (0,0), (-1,-1), 'TOP')
    ]))
    elements.append(KeepTogether(desc_table))
    elements.append(Spacer(1, 15))

    for title, content, img_key in [("ROOT CAUSE", root, "root"), ("L2 ANALYSIS", l2, "l2"), ("RESOLUTION", res, "res")]:
        elements.append(Paragraph(f"<b>{title}</b>", styles["Heading2"]))
        elements.append(Paragraph(content if content else "", styles["Normal"]))
        if images and images.get(img_key):
            elements.append(Spacer(1, 10))
            elements.append(Image(images[img_key], width=450, height=250, kind='proportional'))
        elements.append(Spacer(1, 10))

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 9)
        canvas.drawString(40, 20, str(data.get("number")))
        canvas.drawCentredString(width / 2, 20, f"Page {doc.page}")
        canvas.drawRightString(width - 40, 30, str(data.get("priority")))
        canvas.restoreState()

    doc.build(elements, onFirstPage=footer, onLaterPages=footer)
    buffer.seek(0)
    return buffer
