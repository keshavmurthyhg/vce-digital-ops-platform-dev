from reportlab.lib.pagesizes import letter

def pdf_footer(data):
    def footer(canvas, doc):
        width, _ = letter
        canvas.setFont('Helvetica', 9)

        canvas.drawString(40, 20, str(data.get("number")))
        canvas.drawCentredString(width / 2, 20, f"Page {doc.page}")
        canvas.drawRightString(width - 40, 20, str(data.get("priority")))

    return footer


# WORD FOOTER
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def apply_word_footer(doc, data):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.clear()

    tab_stops = footer.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    run = footer.add_run(f"{data.get('number')}\tPage ")

    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.text = "PAGE"

    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')

    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)

    footer.add_run(f"\t{data.get('priority')}")
