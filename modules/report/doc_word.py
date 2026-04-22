from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import re


def clean_text(text):
    if not text:
        return ""
    return re.sub(r"How does the user want.*?\d+", "", str(text), flags=re.I).strip()


from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO


def generate_word_doc(data, root, l2, res, images=None, full_df=None):

    doc = Document()

    # ---------- TITLE ----------
    doc.add_heading("INCIDENT REPORT", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------- MAIN TABLE ----------
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    def add_link(paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def fill(r, c, key, val):
        h = table.rows[r].cells[c]
        v = table.rows[r].cells[c + 1]

        h.text = key.upper()

        p = v.paragraphs[0]

        if key == "Incident":
            add_link(p, f"https://volvoitsm.service-now.com/nav_to.do?uri=incident.do?sysparm_query=number={val}", str(val))
        elif key == "Azure Bug":
            add_link(p, f"https://dev.azure.com/.../{val}", str(val))
        elif key == "PTC Case":
            add_link(p, f"https://support.ptc.com/.../{val}", str(val))
        else:
            p.text = str(val)

    fill(0,0,"Incident",data.get("number"))
    fill(0,2,"Created By",data.get("created_by"))
    fill(1,0,"Azure Bug",data.get("azure_bug"))
    fill(1,2,"Created Date",data.get("created_date"))
    fill(2,0,"PTC Case",data.get("ptc_case"))
    fill(2,2,"Assigned To",data.get("assigned_to"))
    fill(3,0,"Priority",data.get("priority"))
    fill(3,2,"Resolved Date",data.get("resolved_date"))

    # ---------- DESCRIPTION ----------
    doc.add_heading("Details", 1)
    doc.add_paragraph(data.get("short_description", ""))
    doc.add_paragraph(data.get("description", ""))

    # ---------- SECTIONS ----------
    for title, content, img_key in [
        ("ROOT CAUSE", root, "root"),
        ("L2 ANALYSIS", l2, "l2"),
        ("RESOLUTION", res, "res"),
    ]:
        doc.add_heading(title, 1)
        doc.add_paragraph(content or "")

        if images and images.get(img_key):
            doc.add_picture(images[img_key], width=Inches(4.5))

    # ---------- FOOTER ----------
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = f"{data.get('number')}   |   Page   |   {data.get('priority')}"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
