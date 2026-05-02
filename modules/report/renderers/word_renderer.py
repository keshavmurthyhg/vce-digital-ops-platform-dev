from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from modules.report.layout.footer import apply_word_footer
from modules.common.utils.links import apply_word_link
from modules.common.utils.formatters import (
    format_date,
    set_cell_bg,
    safe_text
)
from modules.common.utils.text_cleaner import (
    clean_text,
    format_description
)


# -----------------------------------
# CELL PADDING
# -----------------------------------
def set_cell_padding(
    cell,
    top=120,
    start=120,
    bottom=120,
    end=120
):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = OxmlElement("w:tcMar")

    for margin_name, margin_value in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        node = OxmlElement(f"w:{margin_name}")
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)

    tcPr.append(tcMar)


def apply_table_padding(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_padding(cell)


def generate_word_doc(
    data,
    root,
    l2,
    res,
    images,
    ppt_data=None
):
    doc = Document()

    # -----------------------------------
    # TITLE
    # -----------------------------------
    title = doc.add_heading(
        "INCIDENT REPORT",
        0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -----------------------------------
    # HEADER TABLE
    # Match PDF ratio: 100:160:100:160
    # Total width ≈ 6.5 inches
    # -----------------------------------
    table = doc.add_table(
        rows=4,
        cols=4
    )

    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    
    # Better proportional sizing
    column_widths = [
        Inches(1.1),   # label
        Inches(2.0),   # value
        Inches(1.4),   # label
        Inches(2.0)    # value
    ]

    for row in table.rows:
        for i, width in enumerate(column_widths):
            row.cells[i].width = width

    def fill(r, c, key, val):
        header_cell = table.rows[r].cells[c]
        value_cell = table.rows[r].cells[c + 1]

        # Header styling
        p = header_cell.paragraphs[0]
        run = p.add_run(key.upper())
        run.bold = True
        set_cell_bg(header_cell)

        cleaned_val = safe_text(val)

        if str(cleaned_val).lower() in [
            "nan",
            "nat",
            "none",
            ""
        ]:
            cleaned_val = "-"

        value_para = value_cell.paragraphs[0]

        apply_word_link(
            value_para,
            key,
            cleaned_val
        )

    fill(0, 0, "Incident", data.get("number"))
    fill(0, 2, "Created By", data.get("created_by"))

    fill(1, 0, "Azure Bug", data.get("azure_bug"))
    fill(
        1,
        2,
        "Created Date",
        format_date(data.get("created_date"))
    )

    fill(2, 0, "PTC Case", data.get("ptc_case"))
    fill(2, 2, "Assigned To", data.get("assigned_to"))

    fill(3, 0, "Priority", data.get("priority"))
    fill(
        3,
        2,
        "Resolved Date",
        format_date(data.get("resolved_date"))
    )

    apply_table_padding(table)

    doc.add_paragraph("")

    # -----------------------------------
    # DESCRIPTION TABLE
    # Same total width as header table
    # -----------------------------------
    t2 = doc.add_table(
        rows=2,
        cols=2
    )
    
    t2.style = "Table Grid"
    t2.autofit = False
    t2.allow_autofit = False
    
    # Proportional widths
    desc_widths = [
        Inches(3.0),   # Short Description
        Inches(3.5)    # Description
    ]
    
    # Apply widths to all rows
    for row in t2.rows:
        for i, width in enumerate(desc_widths):
            row.cells[i].width = width
    
    # Header row
    headers = [
        "SHORT DESCRIPTION",
        "DESCRIPTION"
    ]
    
    for i, text in enumerate(headers):
        cell = t2.rows[0].cells[i]
        p = cell.paragraphs[0]
    
        # Clear default empty paragraph text
        p.clear()
    
        run = p.add_run(text)
        run.bold = True
    
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        set_cell_bg(cell)
    
    # Data row
    t2.rows[1].cells[0].text = clean_text(
        safe_text(
            data.get("short_description")
        )
    )
    
    t2.rows[1].cells[1].text = clean_text(
        format_description(
            safe_text(
                data.get("description")
            )
        )
    )
    
    # Apply padding
    apply_table_padding(t2)
    
    doc.add_paragraph("")

    # -----------------------------------
    # RCA SECTIONS
    # -----------------------------------
    sections = {
        "PROBLEM STATEMENT": root,
        "ROOT CAUSE": l2,
        "RESOLUTION & RECOMMENDATION": res
    }
    
    for title, content in sections.items():
    
        # Section heading
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
        content = safe_text(content)
    
        if str(content).lower() in ["nan", "nat", "none", ""]:
            content = "-"
    
        for line in content.split("\n"):
            cleaned_line = clean_text(
                line.strip("- ").strip()
            )
    
            if cleaned_line:
                p = doc.add_paragraph(
                    cleaned_line,
                    style="List Bullet"
                )
    
                # Match table width alignment
                p.paragraph_format.left_indent = Inches(0.75)
                p.paragraph_format.right_indent = Inches(0.75)
                p.paragraph_format.space_after = Pt(4)
    
        doc.add_paragraph("")
        
    # -----------------------------------
    # FOOTER
    # -----------------------------------
    apply_word_footer(
        doc,
        data
    )

    # -----------------------------------
    # PPT SLIDES
    # -----------------------------------
    if ppt_data:
        try:
            from modules.converter.ppt_slide_renderer import (
                render_ppt_slides_to_images
            )

            doc.add_page_break()
            doc.add_heading(
                "PPT Slides",
                level=1
            )

            slide_images = render_ppt_slides_to_images(
                ppt_data
            )

            if not slide_images:
                doc.add_paragraph(
                    "No slide images found in PPT."
                )
            else:
                for img in slide_images[1:]:
                    doc.add_page_break()
                    doc.add_picture(
                        img,
                        width=Inches(6.5)
                    )

        except Exception as e:
            doc.add_paragraph(
                f"Unable to attach PPT slides: {str(e)}"
            )

    # -----------------------------------
    # SAVE
    # -----------------------------------
    from io import BytesIO

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()
