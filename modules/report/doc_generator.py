from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from datetime import datetime
import re

from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter


# ---------------- COMMON ---------------- #

def clean_text(text):
    if not text:
        return ""
    return re.sub(r"How does the user want.*?\d+", "", str(text), flags=re.I).strip()


def format_date(date_str):
    if not date_str:
        return ""
    try:
        return datetime.strptime(str(date_str), "%Y-%m-%d").strftime("%d-%b-%Y")
    except:
        return str(date_str)


def safe_images(images):
    if not isinstance(images, dict):
        return {"root": [], "l2": [], "res": []}
    return images


# ---------------- WORD ---------------- #

def add_hyperlink(paragraph, url, text):
    if not text:
        paragraph.text = ""
        return

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    text_elem = OxmlElement("w:t")
    text_elem.text = str(text)
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell_bg(cell):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "D9D9D9")
    tcPr.append(shd)


def add_images_word(doc, image_list):
    if not image_list:
        return

    for img in image_list:
        try:
            if hasattr(img, "read"):
                img_bytes = BytesIO(img.read())
                img.seek(0)
            else:
                img_bytes = img

            doc.add_picture(img_bytes, width=Inches(5.5))
            doc.add_paragraph("")
        except:
            pass


def generate_word_doc(data, root, l2, res, images=None):
    images = safe_images(images)

    doc = Document()
    doc.add_heading("INCIDENT REPORT", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    def fill(r, c, key, val, url=None):
        h = table.rows[r].cells[c]
        v = table.rows[r].cells[c + 1]

        p = h.paragraphs[0]
        p.text = key.upper()
        for run in p.runs:
            run.bold = True
        set_cell_bg(h)

        p = v.paragraphs[0]
        if url and val:
            add_hyperlink(p, url, val)
        else:
            p.text = str(val or "")

    fill(0,0,"Incident",data.get("number"),
         f"https://volvoitsm.service-now.com/nav_to.do?uri=incident.do?sysparm_query=number={data.get('number')}")

    fill(0,2,"Created By",data.get("created_by"))
    fill(1,0,"Azure Bug",data.get("azure_bug"),
         f"https://dev.azure.com/VolvoGroup-DVP/VCEWindchillPLM/_workitems/edit/{data.get('azure_bug')}" if data.get("azure_bug") else None)

    fill(1,2,"Created Date",format_date(data.get("created_date")))
    fill(2,0,"PTC Case",data.get("ptc_case"),
         f"https://support.ptc.com/app/caseviewer/?case={data.get('ptc_case')}" if data.get("ptc_case") else None)

    fill(2,2,"Assigned To",data.get("assigned_to"))
    fill(3,0,"Priority",data.get("priority"))
    fill(3,2,"Resolved Date",format_date(data.get("resolved_date")))

    doc.add_paragraph("")

    # DESCRIPTION TABLE
    t2 = doc.add_table(rows=2, cols=2)
    t2.style = "Table Grid"

    headers = ["SHORT DESCRIPTION", "DESCRIPTION"]
    for i, text in enumerate(headers):
        p = t2.rows[0].cells[i].paragraphs[0]
        p.text = text
        for run in p.runs:
            run.bold = True
        set_cell_bg(t2.rows[0].cells[i])

    t2.rows[1].cells[0].text = clean_text(data.get("short_description"))
    t2.rows[1].cells[1].text = clean_text(data.get("description"))

    for row in t2.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Inches(0.05)

    section_map = {
        "PROBLEM STATEMENT & ROOT CAUSE": (root, images.get("root")),
        "TECHNICAL ANALYSIS": (l2, images.get("l2")),
        "RESOLUTION & RECOMMENDATION": (res, images.get("res")),
    }

    for title, (content, imgs) in section_map.items():
        doc.add_heading(title, 1)

        for line in (content or "-").split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("-"):
                line = line[1:].strip()

            p = doc.add_paragraph(line, style="List Bullet")
            p.paragraph_format.space_after = Inches(0.05)

        add_images_word(doc, imgs)

    # FOOTER
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.clear()

    tab_stops = footer.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    run = footer.add_run(f"{data.get('number')}\tPage ")

    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.text = "PAGE"

    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')

    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)

    footer.add_run(f"\t{data.get('priority')}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ---------------- PDF ---------------- #

def add_images_pdf(elements, image_list):
    if not image_list:
        return

    for img in image_list:
        try:
            if hasattr(img, "read"):
                img_bytes = BytesIO(img.read())
                img.seek(0)
            else:
                img_bytes = img

            elements.append(Image(img_bytes, width=400, height=250))
            elements.append(Spacer(1, 10))
        except:
            pass


def generate_pdf(data, root, l2, res, images=None):
    images = safe_images(images)

    buffer = BytesIO()
    styles = getSampleStyleSheet()
    elements = []

    doc = SimpleDocTemplate(buffer, pagesize=letter)

    elements.append(Paragraph("<b>INCIDENT REPORT</b>", styles["Title"]))
    elements.append(Spacer(1,10))

    def wrap_link(text, url=None):
        if not text:
            return Paragraph("", styles["Normal"])
        if url:
            return Paragraph(f'<link href="{url}">{text}</link>', styles["Normal"])
        return Paragraph(str(text), styles["Normal"])

    # HEADER TABLE
    table = Table([
        ["INCIDENT", wrap_link(data.get("number"),
         f"https://volvoitsm.service-now.com/nav_to.do?uri=incident.do?sysparm_query=number={data.get('number')}"),
         "CREATED BY", wrap_link(data.get("created_by"))],

        ["AZURE BUG", wrap_link(data.get("azure_bug"),
         f"https://dev.azure.com/VolvoGroup-DVP/VCEWindchillPLM/_workitems/edit/{data.get('azure_bug')}") if data.get("azure_bug") else "",
         "CREATED DATE", wrap_link(format_date(data.get("created_date")))],

        ["PTC CASE", wrap_link(data.get("ptc_case"),
         f"https://support.ptc.com/app/caseviewer/?case={data.get('ptc_case')}") if data.get("ptc_case") else "",
         "ASSIGNED TO", wrap_link(data.get("assigned_to"))],

        ["PRIORITY", wrap_link(data.get("priority")),
         "RESOLVED DATE", wrap_link(format_date(data.get("resolved_date")))],
    ], colWidths=[100,160,100,160])

    table.setStyle(TableStyle([
        ('GRID',(0,0),(-1,-1),1,colors.black),
        ('BACKGROUND',(0,0),(0,-1),colors.lightgrey),
        ('BACKGROUND',(2,0),(2,-1),colors.lightgrey),
        ('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),
        ('FONTNAME',(2,0),(2,-1),'Helvetica-Bold'),
        ('VALIGN',(0,0),(-1,-1),'TOP'),
    ]))

    elements.append(table)
    elements.append(Spacer(1,15))

    # DESCRIPTION TABLE
    def wrap(x):
        return Paragraph(str(x or ""), styles["Normal"])

    desc = Table([
        [
            Paragraph("<b>SHORT DESCRIPTION</b>", styles["Normal"]),
            Paragraph("<b>DESCRIPTION</b>", styles["Normal"])
        ],
        [
            wrap(clean_text(data.get("short_description"))),
            wrap(clean_text(data.get("description")))
        ]
    ], colWidths=[260,260])

        desc.setStyle(TableStyle([
            ('GRID',(0,0),(-1,-1),1,colors.black),
            ('BACKGROUND',(0,0),(-1,0),colors.lightgrey),
            ('VALIGN',(0,0),(-1,-1),'TOP'),
        
            # ✅ center header row
            ('ALIGN',(0,0),(-1,0),'CENTER'),
        ]))

    elements.append(desc)
    elements.append(Spacer(1,20))

    def add_bullets(text):
        for line in (text or "-").split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("-"):
                line = line[1:].strip()

            elements.append(Paragraph(f"• {line}", styles["Normal"]))
            elements.append(Spacer(1,4))

    def section(title, content, imgs):
        elements.append(Paragraph(f"<b>{title}</b>", styles["Heading2"]))
        elements.append(Spacer(1,6))
        add_bullets(content)
        elements.append(Spacer(1,10))
        add_images_pdf(elements, imgs)

    section("PROBLEM STATEMENT & ROOT CAUSE", root, images.get("root"))
    section("TECHNICAL ANALYSIS", l2, images.get("l2"))
    section("RESOLUTION & RECOMMENDATION", res, images.get("res"))

    # FOOTER
    def footer(canvas, doc):
        width, _ = letter
        canvas.setFont('Helvetica',9)
        canvas.drawString(40,20,str(data.get("number")))
        canvas.drawCentredString(width/2,20,f"Page {doc.page}")
        canvas.drawRightString(width-40,20,str(data.get("priority")))

    doc.build(elements, onFirstPage=footer, onLaterPages=footer)

    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes
