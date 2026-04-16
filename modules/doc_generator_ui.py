import streamlit as st
from modules.snow_loader import load_snow_data
from modules.doc_generator import generate_word_doc
from docx import Document
from io import BytesIO

# ================= FETCH FUNCTION =================
def get_incident_from_df(df, incident_number):

    df_copy = df.copy()

    # normalize
    df_copy["number"] = df_copy["number"].astype(str).str.strip().str.upper()
    incident_number = incident_number.strip().upper()

    row = df_copy[df_copy["number"] == incident_number]

    if not row.empty:
        row = row.iloc[0]

        return {
            "number": row.get("number", ""),
            "short_description": row.get("short description", ""),
            "description": row.get("description", ""),

            "priority": row.get("priority", ""),
            "created_by": row.get("created by", ""),
            "created_date": row.get("created date", ""),
            "assigned_to": row.get("assigned to", ""),
            "resolved_date": row.get("resolved date", ""),

            "work_notes": row.get("work notes", ""),
            "comments": row.get("additional comments", ""),
            "resolution": row.get("resolution notes", ""),

            # OPTIONAL FIELDS (if exist)
           # "azure_bug": row.get(" ", ""),
            "ptc_case": row.get("vendor ticket", "")
        }

    return None
# ================= UI =================
def render_doc_generator():

    st.title("📄 SNOW Incident Report Generator")

    df = load_snow_data()

    # 🔍 DEBUG START (ADD HERE)
    #st.write("Columns:", df.columns)
    #st.write("Sample Row:", df.head(1))
    # 🔍 DEBUG END

    incident_number = st.text_input("Enter Incident Number")

    col1, col2 = st.columns(2)

    # ================= FETCH =================
    if col1.button("Fetch Incident"):

        data = get_incident_from_df(df, incident_number)

        if data:
            st.session_state["doc_data"] = data

            # ✅ FORCE UPDATE TEXT FIELDS
            st.session_state["root"] = data.get("work_notes", "")
            st.session_state["l2"] = data.get("comments", "")
            st.session_state["res"] = data.get("resolution", "")
            st.session_state["closure"] = data.get("resolution", "")

        else:
            st.warning("❌ Incident not found in dataset")

    # ================= FORM =================
    root_cause = st.text_area("Root Cause", key="root")
    l2_analysis = st.text_area("L2 Analysis", key="l2")
    resolution = st.text_area("Resolution", key="res")
    closure = st.text_area("Closure Notes", key="closure")

    # ================= GENERATE =================
    from docx import Document
    from io import BytesIO


    def generate_word_doc(data, root_cause, l2_analysis, resolution, closure):

    doc = Document()

    doc.add_heading('INCIDENT REPORT', 0)

    # ================= TABLE 1 (LINKS) =================
    table1 = doc.add_table(rows=2, cols=3)
    table1.style = 'Table Grid'

    headers = ["Incident", "Azure Bug", "PTC Case"]
    values = [
        data.get("number", ""),
        data.get("azure_bug", ""),
        data.get("ptc_case", "")
    ]

    for i in range(3):
        table1.rows[0].cells[i].text = headers[i]
        table1.rows[1].cells[i].text = str(values[i])

    # ================= TABLE 2 (DETAILS) =================
    table2 = doc.add_table(rows=2, cols=5)
    table2.style = 'Table Grid'

    headers2 = ["Priority", "Created By", "Created Date", "Assigned To", "Resolved Date"]
    values2 = [
        data.get("priority", ""),
        data.get("created_by", ""),
        str(data.get("created_date", "")),
        data.get("assigned_to", ""),
        str(data.get("resolved_date", ""))
    ]

    for i in range(5):
        table2.rows[0].cells[i].text = headers2[i]
        table2.rows[1].cells[i].text = str(values2[i])

    # ================= TABLE 3 (DESCRIPTION) =================
    table3 = doc.add_table(rows=2, cols=2)
    table3.style = 'Table Grid'

    table3.rows[0].cells[0].text = "Short Description"
    table3.rows[0].cells[1].text = "Description"

    table3.rows[1].cells[0].text = data.get("short_description", "")
    table3.rows[1].cells[1].text = data.get("description", "")

    # ================= TEXT SECTIONS =================
    doc.add_heading('Root Cause', 1)
    doc.add_paragraph(root_cause)

    doc.add_heading('L2 Analysis', 1)
    doc.add_paragraph(l2_analysis)

    doc.add_heading('Resolution', 1)
    doc.add_paragraph(resolution)

    doc.add_heading('Closure Notes', 1)
    doc.add_paragraph(closure)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

# ============= Download Documnent ===========================
        st.download_button(
            "📥 Download Report",
            file,
            f"{st.session_state['doc_data']['number']}.docx"
        )
